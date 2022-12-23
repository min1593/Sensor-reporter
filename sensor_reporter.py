import PySimpleGUI as sg
import pandas as pd
import docx, time, glob
import os
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from plotnine import *

sg.theme('dark grey 9') # Add a touch of color

# All the stuff inside your window.
layout = [[sg.Text('Sensor :'), sg.InputText('', size = (15,1), key = 'sen')],
          [sg.Text('First S/N :'), sg.InputText('', size = (15,1), key = 'sn')],
          [sg.Text('Folder Path for Frequency (頻譜) :')],
          [sg.InputText([], size = (40,1), key = 'show1')], 
          [sg.Input(visible = False, enable_events = True, key = 'in'), sg.FolderBrowse()],
          [sg.Text('Folder Path for Sensitivity (示波) :')],
          [sg.InputText([], size = (40,1), key = 'show2')], 
          [sg.Input(visible = False, enable_events = True, key = 'in2'), sg.FolderBrowse()],
          [sg.Column([[sg.Button('Continue'), sg.Exit()]], justification='r')]]

window = sg.Window('Sensor Reporter', layout, location = (800,600), font='12')


# Event Loop to process "events" and get the "values" of the inputs
while True:
    event, values = window.read()
    window['show1'].update(values['in'])
    window['show2'].update(values['in2'])
    SN1 = values['sn']     # 確保 SN 有填寫
    

    if event == sg.WIN_CLOSED or event == 'Exit': # if user closes window or clicks exit
       break
  
    if event == 'Continue' and SN1 == "":
        sg.popup('請填寫SN')

    if event == 'Continue' and SN1 != "":
        path = values['in']
        csv_files = glob.glob(path + "/*.csv")
        csv = sorted(filter(os.path.isfile, csv_files)) # 檔案排序
        print(csv) # print the filename to make sure the list

        i = 1
        for f in csv:
            df = pd.read_csv(f, engine='python', encoding = 'unicode_escape', sep=',', skiprows=12)
            df['Frequency(MHz)'] = df['Frequency(Hz)'] / 1000000
            #新增一個-40定值, 當作標準線
            datas = df.assign(Requirement=-40)

            if df.index[df['Amplitude(dBm)']<-40].tolist() == []:
                p = (ggplot(datas, aes(x='Frequency(MHz)'))
                # 正常數值會低於 -40, 因此將 scale 改成 -50
                + scale_y_continuous(limits=(0,-50))
                + geom_line(aes(y='Amplitude(dBm)'), color='black') 
                + geom_line(aes(y='Requirement'), color='red')
                + ggtitle('Frequency Test')
                )

                name = "fre_pic" + str(i) + ".png"
                # savepic
                p.save(name)
                i = i + 1

            else:
                sg.popup_error(f + ' data over spec(>-40dBm).')
                exit()

        o = 1
        path1 = values['in2']
        csv_files1 = glob.glob(path1 + "/*.csv")
        csv1 = sorted(filter(os.path.isfile, csv_files1)) # 檔案排序
        print(csv1) # print the filename to make sure the list

        for f1 in csv1:
            #read data - 示波器
            df1 = pd.read_csv(f1, engine='python', encoding = 'unicode_escape', sep=',', skiprows=2)
            #將數值轉換成 mV 單位
            df1['Ampl(mV)'] = df1['Ampl'] * 1000
            #新增序號列
            df1['Time'] = range(1,len(df1)+1) 

            #定義為空值
            def no_labels(values):
                return [""] * len(values)

            p1 = (ggplot(df1, aes(x='Time',y='Ampl(mV)'))
                + geom_line()
                + scale_x_continuous(labels=no_labels)
                + ggtitle('Sensitivity Test')
                )

            name1 = "Sen_pic" + str(o) + ".png"
            # savepic
            p1.save(name1)
            o = o + 1

        doc = docx.Document('INSPECTION_Temp.docx')

        doc.add_page_break()
        
        for m in range (1,int(o)):
            # 設定標題名稱 & 置中 & 粗體 & 大小
            SN = doc.add_paragraph()
            SN.alignment = WD_ALIGN_PARAGRAPH.CENTER
            SN2 = SN1[-4:]
            SN3 = int(SN2) + int(m) - 1 
            SN4 = SN1[:-4]
            SN5 = str(SN4) + str(SN3)
            S = SN.add_run(str(SN5))
            S.font.bold = True
            S.font.size = Pt(24)
            pic = "fre_pic" + str(m) + ".png"
            pic1 = "Sen_pic" + str(m) + ".png"             
            doc.add_picture(pic, width=Inches(5.5))
            doc.add_picture(pic1, width=Inches(5.5))
            m = m + 1

        nowTime = int(time.time()) #取得現在時間
        struct_time = time.localtime(nowTime) #轉換成時間元組
        timeString = time.strftime("%Y%m%d_%H%M%S", struct_time) #將時間元組轉換成想要的字串
        file = values['sen']
        filename = file+'_test_'+timeString+".docx" #檔名設定
        doc.save(filename)
        sg.popup('完成')
        break 

        
window.close()